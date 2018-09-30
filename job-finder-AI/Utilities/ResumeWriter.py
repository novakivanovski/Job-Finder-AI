import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class ResumeWriter:
    def __init__(self, resume):
        self.resume = resume
        self.path = os.path.join('Storage', 'config', 'resume', 'generated_resume.docx')
        self.document = Document()
        self.setup_document_style()

    def setup_document_style(self):
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

    def write(self):
        self.write_header()
        self.write_highlights()
        self.write_education()
        self.write_projects()
        self.write_experience()
        self.document.save(self.path)

    def write_header(self):
        resume_header = self.resume.header
        header_paragraph = self.get_header_paragraph()
        name_run = header_paragraph.add_run(resume_header.full_name + '\n')
        name_run.bold = True
        name_run.font.size = Pt(24)
        header_paragraph.add_run(resume_header.email + '\n')
        header_paragraph.add_run('Phone: ' + resume_header.phone_number + '\n')
        header_paragraph.add_run(resume_header.website + '\n')

    def get_header_paragraph(self):
        header_paragraph = self.document.add_paragraph()
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return header_paragraph

    def write_highlights(self):
        highlights = self.resume.highlights
        highlights_paragraph = self.get_section_paragraph(highlights)
        for item in highlights.items:
            highlights_paragraph.add_run(item + '\n')

    def write_education(self):
        education = self.resume.education
        education_paragraph = self.get_section_paragraph(education)
        degree_run = education_paragraph.add_run(education.degree + '\n')
        degree_run.bold = True
        education_paragraph.add_run(education.alma_mater + '\n')
        education_paragraph.add_run(education.awards + '\n')

    def write_projects(self):
        projects = self.resume.projects
        projects_paragraph = self.get_section_paragraph(projects)
        for item in projects.items:
            projects_paragraph.add_run(item + '\n')

    def write_experience(self):
        experience = self.resume.experience
        experience_paragraph = self.get_section_paragraph(experience)
        for job_title in experience.job_titles:
            job_date = experience.job_titles_to_dates[job_title]
            job_location = experience.job_titles_to_locations[job_title]
            job_tasks = experience.job_titles_to_tasks[job_title]
            job_title_run = experience_paragraph.add_run(job_title + '\n')
            job_title_run.bold = True
            experience_paragraph.add_run(job_location + ' '*10 + job_date + '\n')
            for task in job_tasks:
                experience_paragraph.add_run(task + '\n')

    def get_section_paragraph(self, resume_section):
        section_name = resume_section.name
        section_paragraph = self.document.add_paragraph()
        section_name_run = section_paragraph.add_run(section_name + '\n')
        section_name_run.bold = True
        return section_paragraph







